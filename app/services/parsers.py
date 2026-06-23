import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from app.core.errors import AppError


@dataclass
class DocumentElement:
    # 解析后的最小结构单元，可能是一段正文、标题或表格。
    # 后续切块会保留页码/章节信息，用于回答引用。
    content: str
    page_number: Optional[int] = None
    section_title: Optional[str] = None
    element_type: str = "text"
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    # 不同格式的文件最终都归一成 ParsedDocument，处理链路就不用关心 PDF/DOCX/MD 差异。
    elements: list[DocumentElement]
    page_count: Optional[int] = None
    parser_name: str = "unknown"
    ocr_fallback_used: bool = False


def parse_document(path: Path, extension: str) -> ParsedDocument:
    # 根据扩展名选择解析器；不支持的类型在上传阶段通常已被拦截。
    extension = extension.lower()
    if extension in {".txt"}:
        return parse_text_file(path, markdown=False)
    if extension in {".md", ".markdown"}:
        return parse_text_file(path, markdown=True)
    if extension == ".docx":
        return parse_docx(path)
    if extension == ".pdf":
        return parse_pdf(path)
    raise AppError(415, "unsupported_parser_extension", "No parser for this extension.")


def _decode_text(data: bytes) -> str:
    # 文本文件可能来自 Windows/中文环境，按常见编码顺序尝试解码。
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse_text_file(path: Path, markdown: bool) -> ParsedDocument:
    # TXT/Markdown 是最轻量的解析路径，直接读取字节并拆成元素。
    text = _decode_text(path.read_bytes())
    elements = parse_text_elements(text, markdown=markdown)
    if not elements:
        raise AppError(422, "empty_parsed_document", "Parser produced no text.")
    return ParsedDocument(
        elements=elements,
        page_count=None,
        parser_name="markdown" if markdown else "text",
    )


def parse_text_elements(text: str, markdown: bool = False) -> list[DocumentElement]:
    # Markdown 模式下把标题识别为 section_title，后续 citation 能显示章节位置。
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    elements: list[DocumentElement] = []
    current_section: Optional[str] = None
    buffer: list[str] = []

    def flush_buffer() -> None:
        # 空行触发段落结束；连续非空行会合并成一个 paragraph。
        if not buffer:
            return
        paragraph = "\n".join(buffer).strip()
        buffer.clear()
        if paragraph:
            elements.append(
                DocumentElement(
                    content=paragraph,
                    section_title=current_section,
                    element_type="paragraph",
                )
            )

    for raw_line in normalized.split("\n"):
        line = raw_line.strip()
        if markdown:
            heading = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading:
                flush_buffer()
                current_section = heading.group(2).strip()
                elements.append(
                    DocumentElement(
                        content=current_section,
                        section_title=current_section,
                        element_type="heading",
                        metadata={"heading_level": len(heading.group(1))},
                    )
                )
                continue
        if not line:
            flush_buffer()
            continue
        buffer.append(line)
    flush_buffer()
    return elements


def parse_docx(path: Path) -> ParsedDocument:
    # DOCX 优先用 Docling 获得更丰富的结构；未安装时回退到 python-docx。
    docling = _try_docling(path)
    if docling is not None:
        return docling

    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise AppError(
            500,
            "docx_parser_unavailable",
            "DOCX parsing requires docling or python-docx.",
        ) from exc

    document = DocxDocument(str(path))
    elements: list[DocumentElement] = []
    current_section: Optional[str] = None

    for paragraph in document.paragraphs:
        # python-docx 无法像 Docling 那样完整理解版面，但能保留标题样式和段落文本。
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            current_section = text
            element_type = "heading"
        else:
            element_type = "paragraph"
        elements.append(
            DocumentElement(
                content=text,
                section_title=current_section,
                element_type=element_type,
                metadata={"style": paragraph.style.name if paragraph.style else ""},
            )
        )

    for table_index, table in enumerate(document.tables):
        # 表格按“单元格 | 单元格”转成纯文本，保证仍可被检索。
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        content = "\n".join(row for row in rows if row.strip())
        if content:
            elements.append(
                DocumentElement(
                    content=content,
                    section_title=current_section,
                    element_type="table",
                    metadata={"table_index": table_index},
                )
            )

    if not elements:
        raise AppError(422, "empty_parsed_document", "Parser produced no text.")
    return ParsedDocument(elements=elements, parser_name="python-docx")


def parse_pdf(path: Path) -> ParsedDocument:
    # PDF 先尝试 Docling；如果抽取文本过少，则说明可能是扫描版或解析失败。
    docling = _try_docling(path)
    if docling is not None and sum(len(item.content) for item in docling.elements) >= 50:
        return docling

    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise AppError(
            500, "pdf_parser_unavailable", "PDF parsing requires docling or pypdf."
        ) from exc

    reader = PdfReader(str(path))
    elements: list[DocumentElement] = []
    for index, page in enumerate(reader.pages, start=1):
        # pypdf 只能抽取可复制文本，对扫描版 PDF 通常无能为力。
        text = (page.extract_text() or "").strip()
        if text:
            for paragraph in re.split(r"\n\s*\n", text):
                paragraph = paragraph.strip()
                if paragraph:
                    elements.append(
                        DocumentElement(
                            content=paragraph,
                            page_number=index,
                            element_type="paragraph",
                        )
                    )

    total_chars = sum(len(item.content) for item in elements)
    if total_chars < 50:
        return _parse_pdf_with_ocr_fallback(path, len(reader.pages))
    return ParsedDocument(elements=elements, page_count=len(reader.pages), parser_name="pypdf")


def _try_docling(path: Path) -> Optional[ParsedDocument]:
    # Docling 是可选重依赖，不安装时静默返回 None，让轻量环境也能运行。
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        return None

    try:
        converter = DocumentConverter()
        result = converter.convert(str(path))
        markdown = result.document.export_to_markdown()
    except Exception as exc:  # pragma: no cover - optional dependency behavior
        raise AppError(
            422, "docling_parse_failed", f"Docling failed to parse document: {exc}"
        ) from exc

    elements = parse_text_elements(markdown, markdown=True)
    if not elements:
        raise AppError(422, "empty_parsed_document", "Docling produced no text.")
    return ParsedDocument(elements=elements, page_count=None, parser_name="docling")


def _parse_pdf_with_ocr_fallback(path: Path, page_count: int) -> ParsedDocument:
    # 当前 MVP 只检测 OCR 依赖是否存在，尚未实现 PDF 页面渲染到图片的步骤。
    try:
        from paddleocr import PaddleOCR  # noqa: F401
    except ImportError as exc:
        raise AppError(
            422,
            "ocr_fallback_unavailable",
            "PDF text extraction was insufficient and PaddleOCR is not installed.",
            {"page_count": page_count, "path": str(path)},
        ) from exc
    raise AppError(
        501,
        "ocr_rendering_not_configured",
        "PaddleOCR is installed, but PDF page rendering is not configured in this MVP.",
        {"page_count": page_count, "path": str(path)},
    )
